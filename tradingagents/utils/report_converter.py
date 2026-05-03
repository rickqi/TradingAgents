"""Convert Markdown report directories to professional Word (.docx) documents.

Two public entry-points:

1. ``convert_report_dir_to_docx(report_dir)`` — scan a saved report directory
   (with 1_analysts/, 2_research/, … sub-dirs), build a single formatted
   .docx with cover page, TOC, chapters, and headers/footers.

2. ``convert_md_to_docx(md_path)`` / ``convert_all_md_in_dir(directory)`` —
   lightweight pandoc-based conversion (kept for backward-compat).
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Professional DOCX builder (python-docx, no pandoc needed)
# ---------------------------------------------------------------------------

FONT_BODY = "微软雅黑"
FONT_HEADING = "微软雅黑"
COLOR_TITLE = "1F4E79"
COLOR_H1 = "1F4E79"
COLOR_H2 = "2E75B6"
COLOR_H3 = "404040"
PAGE_MARGIN = Cm(2.54)  # 1 inch

# Report section layout: (sub-dir, chapter title, [(file_stem, section title)])
REPORT_STRUCTURE = [
    ("1_analysts", "第一章 分析师团队报告", [
        ("market", "1.1 市场分析师"),
        ("sentiment", "1.2 舆情分析师"),
        ("news", "1.3 新闻分析师"),
        ("fundamentals", "1.4 基本面分析师"),
    ]),
    ("2_research", "第二章 研究团队辩论", [
        ("bull", "2.1 多方研究员"),
        ("bear", "2.2 空方研究员"),
        ("manager", "2.3 研究经理决策"),
    ]),
    ("3_trading", "第三章 交易建议", [
        ("trader", "3.1 交易员提案"),
    ]),
    ("4_risk", "第四章 风险管理辩论", [
        ("aggressive", "4.1 激进分析师"),
        ("conservative", "4.2 保守分析师"),
        ("neutral", "4.3 中立分析师"),
    ]),
    ("5_portfolio", "第五章 最终投资组合决策", [
        ("decision", "5.1 投资组合经理决策"),
    ]),
]


def _set_font(run, name: str = FONT_BODY, size: int = 10.5,
              bold: bool = False, color: str | None = None):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    # Set East-Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_paragraph(doc, text: str, style=None, font_name=FONT_BODY,
                   font_size=10.5, bold=False, color=None,
                   alignment=None, space_before=0, space_after=6,
                   first_line_indent=None):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    _set_font(run, font_name, font_size, bold, color)
    return p


def _strip_markdown(text: str) -> str:
    """Remove common markdown syntax, preserving readable text."""
    # Remove image syntax
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # Remove links but keep text
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    # Remove bold/italic markers
    text = re.sub(r"\*{1,3}(.*?)\*{1,3}", r"\1", text)
    # Remove heading markers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)
    # Clean up excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_md_tables_and_text(text: str):
    """Parse markdown text, yielding ('text', str) or ('table', rows) tuples."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        # Detect table start
        if "|" in line and i + 1 < len(lines) and re.match(r"^[\s|:-]+$", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                rows = _parse_table_lines(table_lines)
                yield ("table", rows)
            continue
        else:
            if line.strip():
                yield ("text", line)
            i += 1


def _parse_table_lines(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into a list of rows (list of cells)."""
    rows = []
    for i, line in enumerate(lines):
        # Skip separator line
        if re.match(r"^[\s|:-]+$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_md_content(doc, text: str):
    """Add markdown content to document, handling tables and text."""
    text = _strip_markdown(text)
    for kind, data in _parse_md_tables_and_text(text):
        if kind == "table":
            _add_table(doc, data)
        else:
            _add_paragraph(doc, data, space_before=1, space_after=3)


def _add_table(doc, rows: list[list[str]]):
    """Add a table to the document."""
    if not rows or not rows[0]:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(_strip_markdown(cell_text))
            is_header = (ri == 0)
            _set_font(run, FONT_BODY, 9 if not is_header else 9.5,
                      bold=is_header, color="FFFFFF" if is_header else None)


def _add_toc(doc):
    """Add a Table of Contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {})
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = p.add_run()
    fld_char_separate = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    run3._element.append(fld_char_separate)

    run4 = p.add_run('（请右键点击此处 → "更新域" 以生成目录）')
    _set_font(run4, FONT_BODY, 10, color="808080")

    run5 = p.add_run()
    fld_char_end = run5._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run5._element.append(fld_char_end)


def _add_header_footer(section, title: str):
    """Add header and footer to a section."""
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    _set_font(run, FONT_BODY, 8, color="808080")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("第 ")
    _set_font(run, FONT_BODY, 8, color="808080")
    # Page number field
    fld1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run2 = fp.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    run2._element.append(instr)
    run3 = fp.add_run()
    fld2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    run3._element.append(fld2)
    run4 = fp.add_run("1")
    _set_font(run4, FONT_BODY, 8, color="808080")
    run5 = fp.add_run()
    fld3 = run5._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run5._element.append(fld3)
    run6 = fp.add_run(" 页")
    _set_font(run6, FONT_BODY, 8, color="808080")


def convert_report_dir_to_docx(report_dir: str | Path,
                               output_path: str | Path | None = None,
                               ticker: str | None = None,
                               analysis_date: str | None = None) -> str:
    """Build a professional Word report from a report directory tree.

    Args:
        report_dir: Path to the report directory (containing 1_analysts/, etc.)
        output_path: Output .docx path. Defaults to <report_dir>/综合分析报告.docx
        ticker: Ticker symbol for the cover page (auto-detected from dir name if None)
        analysis_date: Analysis date for the cover page (auto-detected if None)

    Returns:
        Path to the generated .docx file.
    """
    report_dir = Path(report_dir)
    if not report_dir.is_dir():
        raise FileNotFoundError(f"Report directory not found: {report_dir}")

    # Auto-detect ticker and date from directory name
    dir_name = report_dir.name
    if ticker is None:
        ticker = dir_name.split("_")[0] if "_" in dir_name else dir_name
    if analysis_date is None:
        m = re.search(r"(\d{8})", dir_name)
        if m:
            d = m.group(1)
            analysis_date = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
        else:
            analysis_date = "N/A"

    if output_path is None:
        output_path = report_dir / "综合分析报告.docx"
    output_path = Path(output_path)

    doc = Document()

    # -- Default style --
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # -- Heading styles --
    for level, (size, color) in enumerate([(22, COLOR_H1), (16, COLOR_H2), (13, COLOR_H3)], 1):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_HEADING
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor.from_string(color)
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
        hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        hs.paragraph_format.space_after = Pt(8)

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = PAGE_MARGIN
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN

    # =========================================================================
    # Cover page
    # =========================================================================
    for _ in range(6):
        doc.add_paragraph()

    _add_paragraph(doc, "TradingAgents",
                   font_size=28, bold=True, color=COLOR_TITLE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=0, space_after=12)
    _add_paragraph(doc, "多Agent量化交易分析报告",
                   font_size=18, bold=True, color=COLOR_H2,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=0, space_after=24)

    doc.add_paragraph()  # spacer

    _add_paragraph(doc, f"标的：{ticker}",
                   font_size=14, bold=False, color="404040",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=0, space_after=8)
    _add_paragraph(doc, f"分析日期：{analysis_date}",
                   font_size=12, bold=False, color="808080",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   space_before=0, space_after=8)

    doc.add_paragraph()

    from datetime import datetime
    _add_paragraph(doc, f"报告生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                   font_size=10, color="808080",
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Page break
    doc.add_page_break()

    # =========================================================================
    # TOC
    # =========================================================================
    _add_paragraph(doc, "目  录", font_size=18, bold=True, color=COLOR_TITLE,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    _add_toc(doc)
    doc.add_page_break()

    # =========================================================================
    # Header / Footer
    # =========================================================================
    _add_header_footer(section, f"TradingAgents 分析报告 — {ticker}")

    # =========================================================================
    # Chapters
    # =========================================================================
    has_content = False
    for subdir, chapter_title, sections in REPORT_STRUCTURE:
        chapter_dir = report_dir / subdir
        if not chapter_dir.is_dir():
            continue

        chapter_has_content = False
        for file_stem, section_title in sections:
            md_file = chapter_dir / f"{file_stem}.md"
            if not md_file.exists():
                continue

            if not chapter_has_content:
                # Chapter heading
                _add_paragraph(doc, chapter_title, style="Heading 1")
                chapter_has_content = True
                has_content = True

            # Section heading
            _add_paragraph(doc, section_title, style="Heading 2")

            # Content
            content = md_file.read_text(encoding="utf-8")
            _add_md_content(doc, content)

            # Small spacer after each section
            doc.add_paragraph()

        if chapter_has_content:
            doc.add_page_break()

    if not has_content:
        _add_paragraph(doc, "（未找到任何报告内容）",
                       font_size=12, color="808080",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # =========================================================================
    # Disclaimer
    # =========================================================================
    _add_paragraph(doc, "", space_before=12)
    _add_paragraph(doc, "免责声明",
                   font_size=11, bold=True, color="808080",
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    _add_paragraph(doc,
                   "本报告由 TradingAgents 多Agent量化交易框架自动生成，"
                   "仅供研究参考，不构成任何投资建议。股市有风险，投资需谨慎。",
                   font_size=9, color="808080",
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=0)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Word report saved: %s", output_path)
    return str(output_path)


# ---------------------------------------------------------------------------
# Legacy pandoc-based conversion (kept for backward-compat)
# ---------------------------------------------------------------------------

def _pandoc_available() -> bool:
    return shutil.which("pandoc") is not None


def convert_md_to_docx(md_path: str, docx_path: str = None) -> str | None:
    """Convert a single Markdown file to Word format using pandoc.

    Falls back to python-docx if pandoc is not available.
    """
    md = Path(md_path)
    if not md.exists():
        logger.warning("MD file not found: %s", md)
        return None
    out = Path(docx_path) if docx_path else md.with_suffix(".docx")

    if _pandoc_available():
        try:
            subprocess.run(
                ["pandoc", str(md), "-o", str(out), "--from=markdown", "--to=docx"],
                check=True, capture_output=True, timeout=30,
            )
            logger.info("Converted %s → %s (pandoc)", md.name, out.name)
            return str(out)
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("pandoc failed: %s — falling back to python-docx", exc)

    # Fallback: python-docx simple conversion
    try:
        doc = Document()
        content = md.read_text(encoding="utf-8")
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(out))
        logger.info("Converted %s → %s (python-docx)", md.name, out.name)
        return str(out)
    except Exception as exc:
        logger.warning("python-docx fallback also failed: %s", exc)
        return None


def convert_all_md_in_dir(directory: str) -> list[str]:
    """Convert every .md file in *directory* to .docx (best-effort)."""
    results = []
    for md_file in Path(directory).glob("*.md"):
        result = convert_md_to_docx(str(md_file))
        if result:
            results.append(result)
    return results
